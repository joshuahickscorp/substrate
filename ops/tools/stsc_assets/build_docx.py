"""Build the render-verified STSC-1 construction memo fixture."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 104, 116)
LIGHT = "F2F4F7"


def set_font(run, *, size: float, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margin(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def paragraph_rule(paragraph, color: str = "0B2545") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, INK),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("PROJECT AURORA  /  INTERNAL DECISION MEMO")
    set_font(run, size=9, bold=True, color=MUTED)
    paragraph_rule(header, "B8BCC4")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("STSC-1 construction fixture  •  2026-07-30")
    set_font(run, size=9, color=MUTED)


def add_metadata(document: Document) -> None:
    for label, value in (
        ("To", "Aurora project team"),
        ("From", "Operations lead"),
        ("Date", "July 30, 2026"),
        ("Re", "Recovery plan after the July sensor incident"),
        ("Status", "Decision required before the next field window"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(f"{label}: "), size=11, bold=True, color=INK)
        set_font(paragraph.add_run(value), size=11, color=INK)


def add_table(document: Document) -> None:
    rows = [
        ("Owner", "Workstream", "Due", "Status"),
        ("Mina Chen", "Sensor recalibration", "2026-08-02", "In progress"),
        ("Owen Brooks", "Telemetry reconciliation", "2026-08-03", "Blocked"),
        ("Priya Nair", "Customer communication", "2026-08-04", "Draft"),
        ("Theo Grant", "Field restart checklist", "2026-08-05", "Not started"),
    ]
    table = document.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    widths = (1900, 3500, 1700, 2260)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            set_cell_width(cell, widths[column_index])
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_font(
                paragraph.add_run(value),
                size=10,
                bold=row_index == 0,
                color=INK,
            )
            if row_index == 0:
                shade(cell, LIGHT)


def build(output: Path) -> None:
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("DECISION MEMO"), size=23, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(
        subtitle.add_run("Project Aurora — recovery plan and restart authority"),
        size=14,
        color=MUTED,
    )
    add_metadata(document)
    rule = document.add_paragraph()
    paragraph_rule(rule)

    document.add_heading("Decision requested", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run(
        "Approve a staged field restart only after telemetry reconciliation and "
        "sensor recalibration both pass their independent checks. The proposed "
        "restart window is August 6, 2026 at 09:00 ET."
    )

    document.add_heading("Operating constraints", level=2)
    for text in (
        "Keep the red-channel alert threshold at 72 units until the next calibration review.",
        "Treat the north-array clock drift as unresolved; do not interpolate missing intervals.",
        "Escalate any variance above 4.5% to Mina Chen and record the raw packet digest.",
    ):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(text)

    document.add_heading("Recovery workstreams", level=2)
    add_table(document)

    document.add_heading("Incident evidence", level=2)
    document.add_paragraph(
        "At 14:32 ET on July 28, the north array reported a 7.8% divergence "
        "between the device counter and the reconciled telemetry ledger. A "
        "manual restart reduced divergence to 5.1% but did not restore the "
        "calibration baseline. Packet AUR-0728-1432 remains the evidence anchor."
    )

    acceptance = document.add_paragraph()
    acceptance.paragraph_format.space_before = Pt(6)
    set_font(acceptance.add_run("Acceptance test. "), size=11, bold=True, color=BLUE)
    set_font(
        acceptance.add_run(
            "The restart is admissible only when two consecutive ten-minute "
            "windows show variance at or below 4.5%, no missing packet sequence, "
            "and a matching SHA-256 digest in the telemetry ledger."
        ),
        size=11,
        color=INK,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: build_docx.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())
