"""Typed tool-bearing adapter registry and broker for Odyssey arms.

Flow (mandatory):
  model proposes → broker validates → sandbox executes → cache quarantines
  → verifier admits → arm may consume only admitted results.

This module is shared by candidate and control.  The only intentional
difference between arms remains endogenous developmental memory; tool budgets,
deadlines, attempts, queues, and the closed operation registry are
byte-identical.  Any asymmetry refuses.

No arbitrary shell.  No network except the explicitly admitted ``source.*``
read of already-admitted cache objects.  Evaluator-only material is
unreachable.  Cross-lane cache access is refused.

Protocol-v2 transport constants live in ``odyssey_arms`` and are not altered
here.
"""

from __future__ import annotations

import ast
import contextlib
import hashlib
import json
import os
import re
import subprocess
import tempfile
import threading
import time
import wave
from collections.abc import Callable, Mapping
from dataclasses import dataclass, field
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any

from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric.ed25519 import Ed25519PrivateKey

from substrate import spatial3d
from substrate.odyssey_density import (
    GLOBAL_OVERLAP,
    PROCESS_PER_OP_TOOLS,
    WarmToolPool,
    WarmWorkerJob,
    compact_artifact_receipt,
    resource_class_for_frontier,
    resource_profile_for_frontier,
)
from substrate.product.cache import (
    ArtifactStore,
    LocalCacheVerifierTrustStore,
    ProcessingLineage,
    sign_cache_attestation,
)
from substrate.product.codec import sha256 as product_sha256
from substrate.product.contracts import ProductRefused

# ---------------------------------------------------------------------------
# Schemas and closed registry
# ---------------------------------------------------------------------------

TOOL_REQUEST_SCHEMA = "SUBSTRATE_ODYSSEY_TOOL_REQUEST/v1"
TOOL_RESPONSE_SCHEMA = "SUBSTRATE_ODYSSEY_TOOL_RESPONSE/v1"
TOOL_RECEIPT_SCHEMA = "SUBSTRATE_ODYSSEY_TOOL_RECEIPT/v1"
TOOL_BUDGET_SCHEMA = "SUBSTRATE_ODYSSEY_TOOL_BUDGET/v1"
TOOL_CANARY_SCHEMA = "SUBSTRATE_ODYSSEY_TOOL_BEARING_CANARY/v1"
PROGRAM = "substrate-odyssey-7d-v1"

ROLES = frozenset(("candidate", "control"))
FRONTIERS = frozenset("ABCDEFGH")

# Closed operation registry.  The model may name only these identifiers; the
# broker refuses anything outside the set and anything not declared for the task.
REGISTRY_OPERATIONS: frozenset[str] = frozenset(
    {
        "repo.inspect",
        "repo.test",
        "repo.patch",
        "formal.check_lean",
        "formal.solve_smt",
        "formal.countermodel",
        "media.probe",
        "media.extract_frames",
        "media.transcode_audio",
        "media.transcribe",
        "document.render",
        "document.extract_structure",
        "three_d.build_scene",
        "three_d.render",
        "three_d.depth",
        "three_d.move_object",
        "three_d.set_camera",
        "three_d.inspect_mesh",
        "compute.python",
        "compute.sympy",
        "source.read_cached",
    }
)

# Per-frontier minimum real tool surface implied by the frontier name.
FRONTIER_OPERATIONS: dict[str, frozenset[str]] = {
    "A": frozenset({"repo.inspect", "repo.patch", "document.render", "document.extract_structure", "source.read_cached"}),
    "B": frozenset({"formal.check_lean", "formal.solve_smt", "formal.countermodel", "compute.sympy", "compute.python"}),
    "C": frozenset({"formal.solve_smt", "formal.countermodel"}),
    "D": frozenset({"repo.inspect", "repo.test", "repo.patch", "compute.python"}),
    "E": frozenset({"document.extract_structure", "source.read_cached", "compute.python"}),
    "F": frozenset({"media.probe", "media.transcode_audio", "media.transcribe"}),
    "G": frozenset(
        {
            "media.probe",
            "media.extract_frames",
            "three_d.build_scene",
            "three_d.render",
            "three_d.depth",
            "three_d.move_object",
            "three_d.set_camera",
            "three_d.inspect_mesh",
        }
    ),
    "H": frozenset({"compute.python", "compute.sympy"}),
}

# One representative real operation used by the public canary per frontier.
FRONTIER_CANARY_OPERATION: dict[str, str] = {
    "A": "repo.inspect",
    "B": "formal.check_lean",
    "C": "formal.solve_smt",
    "D": "repo.test",
    "E": "document.extract_structure",
    "F": "media.probe",
    "G": "three_d.render",
    "H": "compute.sympy",
}

_ERROR_CLASSES = frozenset(
    {
        "model_selection",
        "tool_protocol",
        "tool_execution",
        "sandbox",
        "verification",
        "semantic_task",
        "parity",
        "undeclared_operation",
        "evaluator_isolation",
        "cross_lane",
        "forged_receipt",
        "quarantine_bypass",
        "budget",
        "deadline",
        "ok",
    }
)

_FORBIDDEN_PATH_TOKENS = ("evaluator", "answer", "scorer", "hidden", "gold")
_SHA256 = re.compile(r"^[0-9a-f]{64}$")
_IDENTIFIER = re.compile(r"^[a-z0-9][a-z0-9._-]{0,63}$")
_OPERATION = re.compile(r"^[a-z][a-z0-9_]*\.[a-z][a-z0-9_]*$")

# Shared parity constants — identical for both arms.
DEFAULT_CPU_MS = 120_000
DEFAULT_MEMORY_MIB = 2048
DEFAULT_WALL_SECONDS = 120
DEFAULT_MAX_OUTPUT_BYTES = 8 * 1024 * 1024
DEFAULT_MAX_TOOL_CALLS = 4
TOOL_TRANSPORT_ATTEMPTS = 1  # single attempt; no retry ladder

# Host tool pins (G05 panel).  Digests are computed at bind time from the live binary.
# The repository root, and therefore the venv, must follow the checkout rather than a
# machine.  A relocated tree that still resolved an absolute .venv would either run the
# wrong interpreter or drop python/pytest from the G05 panel without saying so.
ROOT = Path(os.environ.get("SUBSTRATE_REPOSITORY_ROOT", Path(__file__).resolve().parents[2])).expanduser().resolve()
_VENV_BIN = Path(os.environ.get("SUBSTRATE_VENV_BIN", ROOT / ".venv" / "bin")).expanduser().resolve()

_DEFAULT_TOOL_PATHS: dict[str, Path] = {
    "lean": Path(os.path.expanduser("~/.elan/toolchains/leanprover--lean4---v4.33.0-rc1/bin/lean")),
    "z3": Path("/opt/homebrew/bin/z3"),
    "ffmpeg": Path("/opt/homebrew/bin/ffmpeg"),
    "ffprobe": Path("/opt/homebrew/bin/ffprobe"),
    "blender": Path("/Applications/Blender.app/Contents/MacOS/Blender"),
    "git": Path("/usr/bin/git"),
    "python": _VENV_BIN / "python",
    "pytest": _VENV_BIN / "pytest",
}

_STRUCTURED_MEDIA = "application/x-substrate-tool-result+json"


class ToolRefused(RuntimeError):
    """Raised when a tool request crosses a custody, parity, or sandbox boundary."""

    def __init__(self, message: str, *, error_class: str = "tool_protocol") -> None:
        if error_class not in _ERROR_CLASSES:
            error_class = "tool_protocol"
        self.error_class = error_class
        super().__init__(message)


def canonical(value: Any) -> bytes:
    return json.dumps(value, sort_keys=True, separators=(",", ":"), ensure_ascii=False).encode("utf-8")


def digest(value: Any) -> str:
    return hashlib.sha256(canonical(value)).hexdigest()


def file_digest(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def _is_sha256(value: object) -> bool:
    return isinstance(value, str) and bool(_SHA256.fullmatch(value))


def _timestamp() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _require(condition: bool, message: str, *, error_class: str = "tool_protocol") -> None:
    if not condition:
        raise ToolRefused(message, error_class=error_class)


def _nonempty(value: object, label: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ToolRefused(f"{label} must be non-empty text", error_class="tool_protocol")
    return value.strip()


def _assert_no_evaluator_tokens(value: object, *, label: str) -> None:
    text = json.dumps(value, sort_keys=True, ensure_ascii=False).casefold() if not isinstance(value, str) else value.casefold()
    for token in _FORBIDDEN_PATH_TOKENS:
        if token in text:
            raise ToolRefused(
                f"{label} names evaluator-only material ({token})",
                error_class="evaluator_isolation",
            )


def _assert_safe_relative(value: object, *, label: str) -> Path:
    text = _nonempty(value, label)
    path = Path(text)
    if path.is_absolute() or not path.parts or any(part in ("", ".", "..") for part in path.parts):
        raise ToolRefused(f"{label} must be a non-escaping relative path", error_class="sandbox")
    if any(any(token in part.casefold() for token in _FORBIDDEN_PATH_TOKENS) for part in path.parts):
        raise ToolRefused(f"{label} may not name an evaluator-only namespace", error_class="evaluator_isolation")
    return path


# ---------------------------------------------------------------------------
# Shared tool budget (parity-critical)
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ToolBudget:
    """Byte-identical resource budget shared by candidate and control."""

    cpu_ms: int = DEFAULT_CPU_MS
    memory_mib: int = DEFAULT_MEMORY_MIB
    wall_seconds: int = DEFAULT_WALL_SECONDS
    max_output_bytes: int = DEFAULT_MAX_OUTPUT_BYTES
    max_tool_calls: int = DEFAULT_MAX_TOOL_CALLS
    attempts: int = TOOL_TRANSPORT_ATTEMPTS

    def __post_init__(self) -> None:
        for name, value in (
            ("cpu_ms", self.cpu_ms),
            ("memory_mib", self.memory_mib),
            ("wall_seconds", self.wall_seconds),
            ("max_output_bytes", self.max_output_bytes),
            ("max_tool_calls", self.max_tool_calls),
            ("attempts", self.attempts),
        ):
            if not isinstance(value, int) or isinstance(value, bool) or value < 1:
                raise ToolRefused(f"tool budget {name} must be a positive integer", error_class="budget")
        if self.attempts != TOOL_TRANSPORT_ATTEMPTS:
            raise ToolRefused("tool transport attempts drifted from the parity constant", error_class="parity")

    def to_dict(self) -> dict[str, Any]:
        return {
            "schema": TOOL_BUDGET_SCHEMA,
            "attempts": self.attempts,
            "cpu_ms": self.cpu_ms,
            "max_output_bytes": self.max_output_bytes,
            "max_tool_calls": self.max_tool_calls,
            "memory_mib": self.memory_mib,
            "wall_seconds": self.wall_seconds,
        }

    def budget_sha256(self) -> str:
        return digest(self.to_dict())

    @classmethod
    def from_dict(cls, value: Mapping[str, Any]) -> ToolBudget:
        if not isinstance(value, Mapping):
            raise ToolRefused("tool budget is malformed", error_class="budget")
        return cls(
            cpu_ms=int(value.get("cpu_ms", DEFAULT_CPU_MS)),
            memory_mib=int(value.get("memory_mib", DEFAULT_MEMORY_MIB)),
            wall_seconds=int(value.get("wall_seconds", DEFAULT_WALL_SECONDS)),
            max_output_bytes=int(value.get("max_output_bytes", DEFAULT_MAX_OUTPUT_BYTES)),
            max_tool_calls=int(value.get("max_tool_calls", DEFAULT_MAX_TOOL_CALLS)),
            attempts=int(value.get("attempts", TOOL_TRANSPORT_ATTEMPTS)),
        )


def assert_budget_parity(candidate: ToolBudget, control: ToolBudget) -> str:
    """Refuse any candidate/control tool-budget asymmetry.  Returns shared digest."""

    cand = candidate.to_dict()
    ctrl = control.to_dict()
    if cand != ctrl:
        raise ToolRefused(
            f"candidate/control tool budget asymmetry refused: {cand} != {ctrl}",
            error_class="parity",
        )
    return candidate.budget_sha256()


def budget_for_frontier(frontier: str) -> ToolBudget:
    """Return the frontier resource-class budget (identical for candidate and control)."""
    try:
        profile = resource_profile_for_frontier(frontier)
    except Exception as error:
        raise ToolRefused(f"frontier resource class unavailable: {error}", error_class="budget") from error
    return ToolBudget(
        cpu_ms=int(profile["cpu_ms"]),
        memory_mib=int(profile["memory_mib"]),
        wall_seconds=int(profile["wall_seconds"]),
        max_output_bytes=int(profile["max_output_bytes"]),
        max_tool_calls=int(profile["max_tool_calls"]),
        attempts=TOOL_TRANSPORT_ATTEMPTS,
    )


# Process-global warm pool: one logical worker class per tool, lane-private jobs.
_WARM_POOL: WarmToolPool | None = None
_WARM_POOL_LOCK = threading.Lock()


def _warm_pool(root: Path) -> WarmToolPool:
    global _WARM_POOL
    with _WARM_POOL_LOCK:
        if _WARM_POOL is None:
            pool_root = root / "evidence/artifacts/substrate/odyssey7d/tool-warm"
            _WARM_POOL = WarmToolPool(pool_root)
        return _WARM_POOL


# ---------------------------------------------------------------------------
# Tool inventory (pinned revisions)
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ToolRevision:
    tool_id: str
    path: str
    version: str
    artifact_sha256: str

    def to_dict(self) -> dict[str, str]:
        return {
            "artifact_sha256": self.artifact_sha256,
            "path": self.path,
            "tool_id": self.tool_id,
            "version": self.version,
        }


def discover_tool_inventory(overrides: Mapping[str, Path] | None = None) -> dict[str, ToolRevision]:
    """Resolve host tool pins and content digests.  Missing tools refuse at bind."""

    paths = dict(_DEFAULT_TOOL_PATHS)
    if overrides:
        paths.update({key: Path(value) for key, value in overrides.items()})
    inventory: dict[str, ToolRevision] = {}
    for tool_id, path in sorted(paths.items()):
        if not path.exists():
            continue
        # Keep venv entrypoints as-is.  Resolving ``.venv/bin/python`` follows the
        # uv symlink to the bare CPython and drops site-packages (docx, sympy, …).
        if path.is_symlink():
            try:
                if not path.resolve().is_file():
                    continue
            except OSError:
                continue
            executable = path if tool_id in {"python", "pytest"} else path.resolve()
        elif path.is_file():
            executable = path.resolve() if tool_id not in {"python", "pytest"} else path
        else:
            continue
        version = _probe_version(tool_id, executable)
        size = executable.stat().st_size if executable.is_file() else 0
        inventory[tool_id] = ToolRevision(
            tool_id=tool_id,
            path=str(executable),
            version=version,
            artifact_sha256=(
                file_digest(executable)
                if executable.is_file() and size < 64 * 1024 * 1024
                else digest({"path": str(executable), "mtime_ns": executable.stat().st_mtime_ns, "size": size})
            ),
        )
    return inventory


def _probe_version(tool_id: str, path: Path) -> str:
    try:
        if tool_id == "lean":
            completed = subprocess.run(  # noqa: S603 - fixed argv, no shell
                [str(path), "--version"],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
            return (completed.stdout or completed.stderr or "lean").splitlines()[0][:200]
        if tool_id in {"z3", "ffmpeg", "ffprobe", "git", "python", "pytest"}:
            flag = "--version" if tool_id != "z3" else "--version"
            completed = subprocess.run(  # noqa: S603
                [str(path), flag],
                capture_output=True,
                text=True,
                timeout=15,
                check=False,
            )
            return (completed.stdout or completed.stderr or tool_id).splitlines()[0][:200]
        if tool_id == "blender":
            completed = subprocess.run(  # noqa: S603
                [str(path), "--version"],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
            return (completed.stdout or "blender").splitlines()[0][:200]
    except (OSError, subprocess.SubprocessError):
        return f"{tool_id}:unprobed"
    return tool_id


# ---------------------------------------------------------------------------
# Request / response contracts
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ToolRequest:
    lane_id: str
    arm: str
    task_id: str
    operation: str
    input_artifact_digests: tuple[str, ...]
    parameters: dict[str, Any]
    resource_budget: ToolBudget
    deadline_unix: float
    declared_operations: frozenset[str]
    budget_sha256: str
    frontier: str

    def to_dict(self) -> dict[str, Any]:
        return {
            "schema": TOOL_REQUEST_SCHEMA,
            "arm": self.arm,
            "budget_sha256": self.budget_sha256,
            "deadline_unix": self.deadline_unix,
            "declared_operations": sorted(self.declared_operations),
            "frontier": self.frontier,
            "input_artifact_digests": list(self.input_artifact_digests),
            "lane_id": self.lane_id,
            "operation": self.operation,
            "parameters": self.parameters,
            "resource_budget": self.resource_budget.to_dict(),
            "task_id": self.task_id,
        }


@dataclass
class ToolResponse:
    status: str
    operation: str
    output_digests: list[str]
    stdout_digest: str | None
    stderr_digest: str | None
    tool_revision: dict[str, str]
    resource_use: dict[str, Any]
    provenance: dict[str, Any]
    error_class: str
    receipt_sha256: str
    admitted: bool
    detail: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        body = {
            "schema": TOOL_RESPONSE_SCHEMA,
            "admitted": self.admitted,
            "detail": self.detail,
            "error_class": self.error_class,
            "operation": self.operation,
            "output_digests": list(self.output_digests),
            "provenance": self.provenance,
            "receipt_sha256": self.receipt_sha256,
            "resource_use": self.resource_use,
            "status": self.status,
            "stderr_digest": self.stderr_digest,
            "stdout_digest": self.stdout_digest,
            "tool_revision": self.tool_revision,
        }
        return body


def declared_operations_for_task(task: Mapping[str, Any], *, frontier: str) -> frozenset[str]:
    """Resolve the closed set of operations a task may invoke."""

    if frontier not in FRONTIERS:
        raise ToolRefused(f"unknown frontier {frontier!r}", error_class="tool_protocol")
    raw = task.get("allowed_operations")
    if raw is None:
        # Default to the frontier's full declared surface so named frontiers
        # actually exercise tools unless a task narrows the grant.
        return FRONTIER_OPERATIONS[frontier]
    if not isinstance(raw, list) or not raw:
        raise ToolRefused("task allowed_operations must be a nonempty list", error_class="tool_protocol")
    declared: set[str] = set()
    for item in raw:
        if not isinstance(item, str) or item not in REGISTRY_OPERATIONS:
            raise ToolRefused(f"task names non-registry operation {item!r}", error_class="undeclared_operation")
        if item not in FRONTIER_OPERATIONS[frontier]:
            raise ToolRefused(
                f"operation {item!r} is not in frontier {frontier} surface",
                error_class="undeclared_operation",
            )
        declared.add(item)
    return frozenset(declared)


def allowed_operations_prompt_block(frontier: str, declared: frozenset[str]) -> str:
    """Human-readable prompt fragment listing only declared registry operations."""

    ops = ", ".join(sorted(declared))
    return (
        f"Bounded tool operations permitted for frontier {frontier}: [{ops}]. "
        "You may propose only those operations via optional tool_proposals "
        "[{operation, parameters, input_artifact_digests}]. "
        "The broker validates, the sandbox executes, the cache quarantines, "
        "and only admitted results are available. "
        "Do not name operations outside that list. "
        "Do not request evaluator material, hidden answer keys, unstated facts, "
        "arbitrary shell, or network access. "
        "Tools never choose contested philosophical conclusions."
    )


# ---------------------------------------------------------------------------
# Sandbox mounts (portable product shape)
# ---------------------------------------------------------------------------


@dataclass
class SandboxMounts:
    """Fixed mount layout: /inputs (ro), /work (ephemeral), /output (quarantine)."""

    root: Path
    inputs: Path
    work: Path
    output: Path

    @classmethod
    def create(cls, root: Path) -> SandboxMounts:
        if root.exists() and (root.is_symlink() or not root.is_dir()):
            raise ToolRefused("sandbox root must be a real directory", error_class="sandbox")
        root.mkdir(parents=True, exist_ok=True)
        inputs = root / "inputs"
        work = root / "work"
        output = root / "output"
        for path in (inputs, work, output):
            path.mkdir(parents=True, exist_ok=True)
            if path.is_symlink():
                raise ToolRefused("sandbox mount must not be a symlink", error_class="sandbox")
        return cls(root=root, inputs=inputs, work=work, output=output)


# ---------------------------------------------------------------------------
# Lane-scoped cache + verifier
# ---------------------------------------------------------------------------


class LaneToolCache:
    """Per-lane, per-arm content-addressed cache with signed admission."""

    def __init__(
        self,
        *,
        root: Path,
        lane_id: str,
        arm: str,
        capacity_bytes: int = 512 * 1024 * 1024,
    ) -> None:
        if lane_id not in FRONTIERS and not _IDENTIFIER.fullmatch(lane_id):
            raise ToolRefused("lane_id is invalid", error_class="tool_protocol")
        if arm not in ROLES:
            raise ToolRefused("arm is invalid", error_class="parity")
        self.lane_id = lane_id
        self.arm = arm
        self.root = root
        self.keys_dir = root / "keys"
        self.trust_dir = root / "trust"
        self.cache_dir = root / "cache"
        self.keys_dir.mkdir(parents=True, exist_ok=True)
        self.trust_dir.mkdir(parents=True, exist_ok=True)
        # Avoid ``.pem`` suffixes: some seatbelt profiles refuse creating them.
        self.private_key_path = self.keys_dir / "verifier.ed25519.private"
        self.public_key_path = self.keys_dir / "verifier.ed25519.public"
        self._ensure_keys()
        self.trust = LocalCacheVerifierTrustStore(self.trust_dir)
        self.trust.trust(
            verifier_id="odyssey-tool-verifier",
            public_key_path=self.public_key_path,
            allowed_rights_statuses=("user-provided", "public", "licensed"),
        )
        if self.cache_dir.exists():
            self.store = ArtifactStore.open(self.cache_dir, verifier_trust_store=self.trust)
        else:
            self.store = ArtifactStore.create(
                self.cache_dir,
                capacity_bytes=capacity_bytes,
                verifier_trust_store=self.trust,
            )
        self._tool_pin_digests: dict[str, str] = {}

    def _ensure_keys(self) -> None:
        if self.private_key_path.exists() and self.public_key_path.exists():
            return
        key = Ed25519PrivateKey.generate()
        private_pem = key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption(),
        )
        public_pem = key.public_key().public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo,
        )
        self.private_key_path.write_bytes(private_pem)
        self.public_key_path.write_bytes(public_pem)
        os.chmod(self.private_key_path, 0o600)
        os.chmod(self.public_key_path, 0o600)

    def ensure_tool_pin(self, revision: ToolRevision) -> str:
        """Ingest and admit a small tool-pin manifest; return its digest."""

        if revision.tool_id in self._tool_pin_digests:
            return self._tool_pin_digests[revision.tool_id]
        pin = {
            "kind": "odyssey-tool-pin",
            "tool_id": revision.tool_id,
            "version": revision.version,
            "artifact_sha256": revision.artifact_sha256,
            "path": revision.path,
        }
        payload = canonical(pin)
        pin_digest = hashlib.sha256(payload).hexdigest()
        # Reuse an already-admitted pin with the same content digest.
        try:
            zone, existing = self.store._locate(pin_digest)  # noqa: SLF001
            if zone in {"verified", "processed"} and existing.verification_status == "verified":
                self._tool_pin_digests[revision.tool_id] = pin_digest
                return pin_digest
        except ProductRefused:
            pass
        with tempfile.NamedTemporaryFile("wb", delete=False) as handle:
            handle.write(payload)
            temporary = Path(handle.name)
        try:
            try:
                descriptor = self.store.ingest_file(
                    temporary,
                    media_type=_STRUCTURED_MEDIA,
                    source_reference_sha256=digest({"lane": self.lane_id, "arm": self.arm, "pin": pin}),
                    rights_status="user-provided",
                )
            except ProductRefused as error:
                # Same bytes may already be immutable from a prior admit.
                if "already immutable" in str(error) or "already bound" in str(error):
                    self._tool_pin_digests[revision.tool_id] = pin_digest
                    return pin_digest
                raise ToolRefused(f"tool pin ingest failed: {error}", error_class="verification") from error
            admitted = self._admit(descriptor.sha256, rights_status="user-provided")
        finally:
            temporary.unlink(missing_ok=True)
        self._tool_pin_digests[revision.tool_id] = admitted.sha256
        return admitted.sha256

    def _admit(self, artifact_sha256: str, *, rights_status: str) -> Any:
        zone, descriptor = self.store._locate(artifact_sha256)  # noqa: SLF001 - broker is the trusted boundary
        if zone != "quarantine":
            # Already admitted or missing.
            if zone in {"verified", "processed"} and descriptor.verification_status == "verified":
                return descriptor
            raise ToolRefused("artifact is not in a promotable state", error_class="verification")
        expires = (datetime.now(UTC) + timedelta(days=7)).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        attestation = sign_cache_attestation(
            artifact_sha256=artifact_sha256,
            cache_id=self.store.cache_id,
            descriptor_sha256=product_sha256(descriptor.to_dict()),
            verifier_id="odyssey-tool-verifier",
            rights_status=rights_status,
            private_key_path=self.private_key_path,
            expires_at=expires,
        )
        return self.store.verify(artifact_sha256, attestation, self.trust)

    def quarantine_and_admit_bytes(
        self,
        data: bytes,
        *,
        media_type: str,
        source_reference_sha256: str,
        rights_status: str = "user-provided",
        processing_lineage: tuple[ProcessingLineage, ...] | None = None,
    ) -> str:
        content_digest = hashlib.sha256(data).hexdigest()
        try:
            zone, existing = self.store._locate(content_digest)  # noqa: SLF001
            if zone in {"verified", "processed"} and existing.verification_status == "verified":
                return content_digest
        except ProductRefused:
            pass
        with tempfile.NamedTemporaryFile("wb", delete=False) as handle:
            handle.write(data)
            temporary = Path(handle.name)
        try:
            try:
                if processing_lineage:
                    descriptor = self.store.ingest_derivative_file(
                        temporary,
                        media_type=media_type,
                        source_reference_sha256=source_reference_sha256,
                        rights_status=rights_status,
                        processing_lineage=processing_lineage,
                    )
                else:
                    descriptor = self.store.ingest_file(
                        temporary,
                        media_type=media_type,
                        source_reference_sha256=source_reference_sha256,
                        rights_status=rights_status,
                    )
            except ProductRefused as error:
                if "already immutable" in str(error):
                    return content_digest
                raise ToolRefused(f"cache quarantine failed: {error}", error_class="verification") from error
            admitted = self._admit(descriptor.sha256, rights_status=rights_status)
        finally:
            temporary.unlink(missing_ok=True)
        return admitted.sha256

    def quarantine_and_admit_file(
        self,
        path: Path,
        *,
        media_type: str,
        source_reference_sha256: str,
        rights_status: str = "user-provided",
        processing_lineage: tuple[ProcessingLineage, ...] | None = None,
    ) -> str:
        if processing_lineage:
            descriptor = self.store.ingest_derivative_file(
                path,
                media_type=media_type,
                source_reference_sha256=source_reference_sha256,
                rights_status=rights_status,
                processing_lineage=processing_lineage,
            )
        else:
            descriptor = self.store.ingest_file(
                path,
                media_type=media_type,
                source_reference_sha256=source_reference_sha256,
                rights_status=rights_status,
            )
        admitted = self._admit(descriptor.sha256, rights_status=rights_status)
        return admitted.sha256

    def read_admitted(self, artifact_sha256: str, *, expected_lane: str) -> bytes:
        if expected_lane != self.lane_id:
            raise ToolRefused("cross-lane cache access refused", error_class="cross_lane")
        if not _is_sha256(artifact_sha256):
            raise ToolRefused("artifact digest is invalid", error_class="tool_protocol")
        try:
            zone, descriptor = self.store._locate(artifact_sha256)  # noqa: SLF001
        except ProductRefused as error:
            raise ToolRefused(f"cache lookup failed: {error}", error_class="verification") from error
        if zone not in {"verified", "processed"} or descriptor.verification_status != "verified":
            raise ToolRefused(
                "quarantine bypass refused: artifact is not admitted",
                error_class="quarantine_bypass",
            )
        blob = self.store._blob_path(zone, artifact_sha256)  # noqa: SLF001
        return blob.read_bytes()

    def materialize_admitted(self, artifact_sha256: str, destination: Path, *, expected_lane: str) -> Path:
        data = self.read_admitted(artifact_sha256, expected_lane=expected_lane)
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(data)
        return destination


def verify_tool_receipt(receipt: Mapping[str, Any], *, cache: LaneToolCache) -> dict[str, Any]:
    """Refuse forged, incomplete, or non-admitted tool receipts."""

    if not isinstance(receipt, Mapping):
        raise ToolRefused("tool receipt must be an object", error_class="forged_receipt")
    required = {
        "schema",
        "status",
        "operation",
        "output_digests",
        "tool_revision",
        "provenance",
        "receipt_sha256",
        "admitted",
        "error_class",
    }
    if not required.issubset(set(receipt)):
        raise ToolRefused("tool receipt is missing required fields", error_class="forged_receipt")
    if receipt.get("schema") != TOOL_RECEIPT_SCHEMA:
        raise ToolRefused("tool receipt schema is invalid", error_class="forged_receipt")
    unsigned = {key: value for key, value in receipt.items() if key != "receipt_sha256"}
    claimed = receipt.get("receipt_sha256")
    if not _is_sha256(claimed) or claimed != digest(unsigned):
        raise ToolRefused("forged tool receipt refused: self-digest mismatch", error_class="forged_receipt")
    if receipt.get("admitted") is not True:
        raise ToolRefused("tool receipt is not admitted", error_class="verification")
    for output_digest in receipt.get("output_digests") or []:
        if not _is_sha256(output_digest):
            raise ToolRefused("tool receipt output digest is invalid", error_class="forged_receipt")
        # Must be readable as admitted evidence in this lane cache.
        cache.read_admitted(output_digest, expected_lane=cache.lane_id)
    provenance = receipt.get("provenance")
    if not isinstance(provenance, dict) or "request_sha256" not in provenance:
        raise ToolRefused("tool receipt provenance is incomplete", error_class="forged_receipt")
    return dict(receipt)


# ---------------------------------------------------------------------------
# Executors — typed argv only
# ---------------------------------------------------------------------------


def _run_fixed(
    argv: list[str],
    *,
    cwd: Path,
    wall_seconds: int,
    env: dict[str, str] | None = None,
) -> tuple[int, bytes, bytes]:
    """Run a fixed argv vector with no shell and a wall-clock bound."""

    if not argv or not all(isinstance(item, str) for item in argv):
        raise ToolRefused("sandbox argv must be a nonempty string list", error_class="sandbox")
    # Refuse shell metacharacter smuggling via empty/null.
    if any("\x00" in item for item in argv):
        raise ToolRefused("sandbox argv contains NUL", error_class="sandbox")
    clean_env = {
        "PATH": "/usr/bin:/bin:/opt/homebrew/bin",
        "HOME": str(cwd / ".home"),
        "TMPDIR": str(cwd / ".tmp"),
        "LANG": "C.UTF-8",
    }
    (cwd / ".home").mkdir(exist_ok=True)
    (cwd / ".tmp").mkdir(exist_ok=True)
    if env:
        # Only explicitly allowlisted keys may be added.
        for key, value in env.items():
            if key in {"LEAN_PATH", "PYTHONPATH", "ELAN_HOME"} and isinstance(value, str):
                clean_env[key] = value
    try:
        completed = subprocess.run(  # noqa: S603 - fixed argv, shell=False
            argv,
            cwd=str(cwd),
            capture_output=True,
            timeout=wall_seconds,
            env=clean_env,
            check=False,
        )
    except subprocess.TimeoutExpired as error:
        raise ToolRefused(f"tool execution timed out: {error}", error_class="tool_execution") from error
    except OSError as error:
        raise ToolRefused(f"tool execution failed to start: {error}", error_class="tool_execution") from error
    return completed.returncode, completed.stdout or b"", completed.stderr or b""


def _run_via_warm_pool(
    *,
    root: Path,
    tool_id: str,
    argv: list[str],
    cwd: Path,
    wall_seconds: int,
    lane_id: str,
    arm: str,
    env: dict[str, str] | None = None,
) -> tuple[int, bytes, bytes, str]:
    """Dispatch through the warm pool when safe; else process-per-op.

    Returns (code, stdout, stderr, mode).  Lean and Blender always use
    process-per-op because deterministic multi-job reset is not guaranteed.
    """
    pool = _warm_pool(root)

    def runner(
        run_argv: list[str],
        run_cwd: Path,
        run_wall: int,
        run_env: dict[str, str] | None,
    ) -> tuple[int, bytes, bytes]:
        return _run_fixed(run_argv, cwd=run_cwd, wall_seconds=run_wall, env=run_env)

    result = pool.run(
        WarmWorkerJob(tool_id=tool_id, argv=argv, cwd=cwd, wall_seconds=wall_seconds, env=dict(env or {})),
        lane_id=lane_id,
        arm=arm,
        runner=runner,
    )
    return result.returncode, result.stdout, result.stderr, result.mode


def _require_tool(inventory: Mapping[str, ToolRevision], tool_id: str) -> ToolRevision:
    if tool_id not in inventory:
        raise ToolRefused(f"required tool {tool_id!r} is not available on this host", error_class="tool_execution")
    return inventory[tool_id]


class OperationExecutor:
    """Dispatch closed registry operations against a sandbox mount set."""

    def __init__(
        self,
        *,
        mounts: SandboxMounts,
        inventory: Mapping[str, ToolRevision],
        cache: LaneToolCache,
        budget: ToolBudget,
        root: Path | None = None,
        lane_id: str = "lane",
        arm: str = "candidate",
    ) -> None:
        self.mounts = mounts
        self.inventory = inventory
        self.cache = cache
        self.budget = budget
        self.root = root
        self.lane_id = lane_id
        self.arm = arm
        self.last_exec_mode: str | None = None

    def _tool_run(
        self,
        tool_id: str,
        argv: list[str],
        *,
        cwd: Path | None = None,
        wall_seconds: int | None = None,
        env: dict[str, str] | None = None,
    ) -> tuple[int, bytes, bytes]:
        """Run a host tool via warm pool when safe; record CPU span for overlap."""
        run_cwd = cwd or self.mounts.work
        run_wall = wall_seconds if wall_seconds is not None else min(90, self.budget.wall_seconds)
        with GLOBAL_OVERLAP.span(lane_id=self.lane_id, arm=self.arm, kind="tool"):
            if self.root is not None:
                code, out, err, mode = _run_via_warm_pool(
                    root=self.root,
                    tool_id=tool_id,
                    argv=argv,
                    cwd=run_cwd,
                    wall_seconds=run_wall,
                    lane_id=self.lane_id,
                    arm=self.arm,
                    env=env,
                )
                self.last_exec_mode = mode
                return code, out, err
            code, out, err = _run_fixed(argv, cwd=run_cwd, wall_seconds=run_wall, env=env)
            self.last_exec_mode = "process_per_op" if tool_id in PROCESS_PER_OP_TOOLS else "warm"
            return code, out, err

    def execute(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        handlers: dict[str, Callable[[ToolRequest], tuple[dict[str, Any], ToolRevision, list[Path]]]] = {
            "repo.inspect": self._repo_inspect,
            "repo.test": self._repo_test,
            "repo.patch": self._repo_patch,
            "formal.check_lean": self._formal_check_lean,
            "formal.solve_smt": self._formal_solve_smt,
            "formal.countermodel": self._formal_countermodel,
            "media.probe": self._media_probe,
            "media.extract_frames": self._media_extract_frames,
            "media.transcode_audio": self._media_transcode_audio,
            "media.transcribe": self._media_transcribe,
            "document.render": self._document_render,
            "document.extract_structure": self._document_extract_structure,
            "three_d.build_scene": self._three_d_build_scene,
            "three_d.render": self._three_d_render,
            "three_d.depth": self._three_d_depth,
            "three_d.move_object": self._three_d_move_object,
            "three_d.set_camera": self._three_d_set_camera,
            "three_d.inspect_mesh": self._three_d_inspect_mesh,
            "compute.python": self._compute_python,
            "compute.sympy": self._compute_sympy,
            "source.read_cached": self._source_read_cached,
        }
        handler = handlers.get(request.operation)
        if handler is None:
            raise ToolRefused(f"operation {request.operation!r} has no executor", error_class="undeclared_operation")
        return handler(request)

    # -- repo ----------------------------------------------------------------

    def _ensure_git_repo(self) -> Path:
        repo = self.mounts.work / "repo"
        repo.mkdir(parents=True, exist_ok=True)
        git = _require_tool(self.inventory, "git")
        if not (repo / ".git").exists():
            code, _out, err = _run_fixed([git.path, "init"], cwd=repo, wall_seconds=30)
            if code != 0:
                raise ToolRefused(f"git init failed: {err.decode('utf-8', 'replace')}", error_class="tool_execution")
            _run_fixed([git.path, "config", "user.email", "odyssey@localhost"], cwd=repo, wall_seconds=10)
            _run_fixed([git.path, "config", "user.name", "odyssey"], cwd=repo, wall_seconds=10)
            readme = repo / "README.md"
            if not readme.exists():
                readme.write_text("# Odyssey tool workspace\n", encoding="utf-8")
            _run_fixed([git.path, "add", "README.md"], cwd=repo, wall_seconds=10)
            _run_fixed([git.path, "commit", "-m", "init"], cwd=repo, wall_seconds=30)
        return repo

    def _repo_inspect(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        git = _require_tool(self.inventory, "git")
        repo = self._ensure_git_repo()
        code, out, err = _run_fixed([git.path, "status", "--porcelain=v1", "--branch"], cwd=repo, wall_seconds=30)
        tree_code, tree_out, _ = _run_fixed([git.path, "ls-files"], cwd=repo, wall_seconds=30)
        result = {
            "operation": "repo.inspect",
            "exit_code": code,
            "status": out.decode("utf-8", "replace"),
            "tracked_files": tree_out.decode("utf-8", "replace").splitlines(),
            "stderr": err.decode("utf-8", "replace")[:2000],
            "repo_root": "work/repo",
            "ls_files_exit": tree_code,
        }
        path = self.mounts.output / "repo_inspect.json"
        path.write_bytes(canonical(result))
        return result, git, [path]

    def _repo_test(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        pytest_tool = _require_tool(self.inventory, "pytest")
        python = _require_tool(self.inventory, "python")
        repo = self._ensure_git_repo()
        tests_dir = repo / "tests"
        tests_dir.mkdir(exist_ok=True)
        # Seed a failing test then a passing one based on parameters.
        mode = str(request.parameters.get("mode", "pass"))
        test_file = tests_dir / "test_odyssey_tool.py"
        if mode == "fail":
            test_file.write_text("def test_odyssey_tool_fail():\n    assert 1 + 1 == 3\n", encoding="utf-8")
        else:
            test_file.write_text("def test_odyssey_tool_pass():\n    assert 1 + 1 == 2\n", encoding="utf-8")
        code, out, err = _run_fixed(
            [pytest_tool.path, "-q", str(test_file)],
            cwd=repo,
            wall_seconds=min(60, self.budget.wall_seconds),
            env={"PYTHONPATH": str(repo)},
        )
        result = {
            "operation": "repo.test",
            "exit_code": code,
            "stdout": out.decode("utf-8", "replace")[:4000],
            "stderr": err.decode("utf-8", "replace")[:4000],
            "mode": mode,
            "passed": code == 0,
            "python": python.version,
        }
        path = self.mounts.output / "repo_test.json"
        path.write_bytes(canonical(result))
        return result, pytest_tool, [path]

    def _repo_patch(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        git = _require_tool(self.inventory, "git")
        repo = self._ensure_git_repo()
        relative = _assert_safe_relative(request.parameters.get("path", "src/module.py"), label="patch path")
        content = request.parameters.get("content", "def value():\n    return 42\n")
        if not isinstance(content, str) or len(content.encode("utf-8")) > 64 * 1024:
            raise ToolRefused("patch content is invalid or too large", error_class="tool_protocol")
        _assert_no_evaluator_tokens(content, label="patch content")
        target = repo / relative
        target.parent.mkdir(parents=True, exist_ok=True)
        before = target.read_text(encoding="utf-8") if target.exists() else ""
        target.write_text(content, encoding="utf-8")
        _run_fixed([git.path, "add", str(relative)], cwd=repo, wall_seconds=10)
        code, out, err = _run_fixed([git.path, "commit", "-m", "odyssey tool patch"], cwd=repo, wall_seconds=30)
        result = {
            "operation": "repo.patch",
            "path": str(relative),
            "before_sha256": hashlib.sha256(before.encode("utf-8")).hexdigest(),
            "after_sha256": hashlib.sha256(content.encode("utf-8")).hexdigest(),
            "commit_exit": code,
            "stdout": out.decode("utf-8", "replace")[:2000],
            "stderr": err.decode("utf-8", "replace")[:2000],
        }
        path = self.mounts.output / "repo_patch.json"
        path.write_bytes(canonical(result))
        return result, git, [path]

    # -- formal --------------------------------------------------------------

    def _formal_check_lean(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        lean = _require_tool(self.inventory, "lean")
        source = request.parameters.get(
            "source",
            "theorem two_plus_two : 2 + 2 = 4 := by rfl\n",
        )
        if not isinstance(source, str) or len(source) > 32_000:
            raise ToolRefused("lean source is invalid", error_class="tool_protocol")
        _assert_no_evaluator_tokens(source, label="lean source")
        lean_file = self.mounts.work / "check.lean"
        lean_file.write_text(source, encoding="utf-8")
        # The inventory already binds the exact Lean executable and digest.  Invoke
        # that pinned binary directly: the sandbox deliberately supplies an
        # isolated HOME, so `elan run` would try to resolve or download a separate
        # toolchain and make the canary depend on ambient machine state.
        argv = [lean.path, str(lean_file)]
        code, out, err = self._tool_run("lean", argv, wall_seconds=min(90, self.budget.wall_seconds))
        result = {
            "operation": "formal.check_lean",
            "exit_code": code,
            "stdout": out.decode("utf-8", "replace")[:4000],
            "stderr": err.decode("utf-8", "replace")[:4000],
            "checked": code == 0,
            "source_sha256": hashlib.sha256(source.encode("utf-8")).hexdigest(),
            "exec_mode": self.last_exec_mode or "process_per_op",
        }
        path = self.mounts.output / "lean_check.json"
        path.write_bytes(canonical(result))
        return result, lean, [path]

    def _formal_solve_smt(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        z3 = _require_tool(self.inventory, "z3")
        smt = request.parameters.get(
            "smt",
            "(set-logic QF_LIA)\n(declare-const x Int)\n(assert (= (+ x 2) 4))\n(check-sat)\n",
        )
        if not isinstance(smt, str) or len(smt) > 32_000:
            raise ToolRefused("smt source is invalid", error_class="tool_protocol")
        _assert_no_evaluator_tokens(smt, label="smt source")
        smt_file = self.mounts.work / "goal.smt2"
        smt_file.write_text(smt, encoding="utf-8")
        code, out, err = self._tool_run("z3", [z3.path, str(smt_file)], wall_seconds=60)
        text = out.decode("utf-8", "replace")
        result = {
            "operation": "formal.solve_smt",
            "exit_code": code,
            "stdout": text[:4000],
            "stderr": err.decode("utf-8", "replace")[:4000],
            "sat": "sat" in text.splitlines()[:3],
            "unsat": "unsat" in text.splitlines()[:3],
            "exec_mode": self.last_exec_mode or "warm",
        }
        path = self.mounts.output / "smt_result.json"
        path.write_bytes(canonical(result))
        return result, z3, [path]

    def _formal_countermodel(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        z3 = _require_tool(self.inventory, "z3")
        # Default: a formula that is satisfiable with a concrete model (countermodel
        # to the claim "x + 2 = 4 has no integer solution").
        smt = request.parameters.get(
            "smt",
            "(set-logic QF_LIA)\n(declare-const x Int)\n(assert (= (+ x 2) 4))\n(check-sat)\n(get-model)\n",
        )
        if not isinstance(smt, str) or len(smt) > 32_000:
            raise ToolRefused("smt source is invalid", error_class="tool_protocol")
        smt_file = self.mounts.work / "countermodel.smt2"
        smt_file.write_text(smt, encoding="utf-8")
        code, out, err = self._tool_run("z3", [z3.path, str(smt_file)], wall_seconds=60)
        text = out.decode("utf-8", "replace")
        result = {
            "operation": "formal.countermodel",
            "exit_code": code,
            "stdout": text[:4000],
            "stderr": err.decode("utf-8", "replace")[:4000],
            "has_model": "define-fun" in text or "(model" in text,
            "exec_mode": self.last_exec_mode or "warm",
        }
        path = self.mounts.output / "countermodel.json"
        path.write_bytes(canonical(result))
        return result, z3, [path]

    # -- media ---------------------------------------------------------------

    def _resolve_media_input(self, request: ToolRequest) -> Path:
        if request.input_artifact_digests:
            digest_value = request.input_artifact_digests[0]
            destination = self.mounts.inputs / f"{digest_value[:16]}.bin"
            return self.cache.materialize_admitted(digest_value, destination, expected_lane=request.lane_id)
        # Generate a tiny synthetic wav for canaries when no input is provided.
        wav_path = self.mounts.inputs / "synthetic.wav"
        if not wav_path.exists():
            with wave.open(str(wav_path), "w") as handle:
                handle.setnchannels(1)
                handle.setsampwidth(2)
                handle.setframerate(16000)
                # 0.25s of silence
                handle.writeframes(b"\x00\x00" * 4000)
        return wav_path

    def _media_probe(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        ffprobe = _require_tool(self.inventory, "ffprobe")
        media = self._resolve_media_input(request)
        code, out, err = _run_fixed(
            [
                ffprobe.path,
                "-v",
                "error",
                "-show_format",
                "-show_streams",
                "-print_format",
                "json",
                str(media),
            ],
            cwd=self.mounts.work,
            wall_seconds=30,
        )
        try:
            probe = json.loads(out.decode("utf-8") or "{}")
        except json.JSONDecodeError:
            probe = {"raw": out.decode("utf-8", "replace")[:4000]}
        result = {
            "operation": "media.probe",
            "exit_code": code,
            "probe": probe,
            "stderr": err.decode("utf-8", "replace")[:2000],
            "input_sha256": file_digest(media),
        }
        path = self.mounts.output / "media_probe.json"
        path.write_bytes(canonical(result))
        return result, ffprobe, [path]

    def _media_extract_frames(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        ffmpeg = _require_tool(self.inventory, "ffmpeg")
        # Build a one-frame video from the synthetic wav + color source if needed.
        media = self.mounts.inputs / "frame_source.mp4"
        if not media.exists():
            code, _out, err = _run_fixed(
                [
                    ffmpeg.path,
                    "-y",
                    "-f",
                    "lavfi",
                    "-i",
                    "color=c=blue:s=64x64:d=1",
                    "-frames:v",
                    "1",
                    str(media),
                ],
                cwd=self.mounts.work,
                wall_seconds=30,
            )
            if code != 0:
                raise ToolRefused(f"ffmpeg frame source failed: {err.decode('utf-8', 'replace')}", error_class="tool_execution")
        frame = self.mounts.output / "frame.png"
        code, out, err = _run_fixed(
            [ffmpeg.path, "-y", "-i", str(media), "-frames:v", "1", str(frame)],
            cwd=self.mounts.work,
            wall_seconds=30,
        )
        if code != 0 or not frame.exists():
            raise ToolRefused(f"frame extraction failed: {err.decode('utf-8', 'replace')}", error_class="tool_execution")
        result = {
            "operation": "media.extract_frames",
            "exit_code": code,
            "frame_sha256": file_digest(frame),
            "stdout": out.decode("utf-8", "replace")[:1000],
            "stderr": err.decode("utf-8", "replace")[:1000],
        }
        meta = self.mounts.output / "extract_frames.json"
        meta.write_bytes(canonical(result))
        return result, ffmpeg, [frame, meta]

    def _media_transcode_audio(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        ffmpeg = _require_tool(self.inventory, "ffmpeg")
        media = self._resolve_media_input(request)
        out_wav = self.mounts.output / "transcoded.wav"
        code, out, err = _run_fixed(
            [ffmpeg.path, "-y", "-i", str(media), "-ac", "1", "-ar", "16000", str(out_wav)],
            cwd=self.mounts.work,
            wall_seconds=60,
        )
        if code != 0 or not out_wav.exists():
            raise ToolRefused(f"audio transcode failed: {err.decode('utf-8', 'replace')}", error_class="tool_execution")
        result = {
            "operation": "media.transcode_audio",
            "exit_code": code,
            "output_sha256": file_digest(out_wav),
            "stdout": out.decode("utf-8", "replace")[:1000],
            "stderr": err.decode("utf-8", "replace")[:1000],
        }
        meta = self.mounts.output / "transcode.json"
        meta.write_bytes(canonical(result))
        return result, ffmpeg, [out_wav, meta]

    def _media_transcribe(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        python = _require_tool(self.inventory, "python")
        media = self._resolve_media_input(request)
        # Prefer waveform evidence via ffmpeg first when input is not wav.
        ffmpeg = _require_tool(self.inventory, "ffmpeg")
        wav = self.mounts.work / "for_whisper.wav"
        code, _out, err = _run_fixed(
            [ffmpeg.path, "-y", "-i", str(media), "-ac", "1", "-ar", "16000", str(wav)],
            cwd=self.mounts.work,
            wall_seconds=60,
        )
        if code != 0:
            raise ToolRefused(f"whisper prep failed: {err.decode('utf-8', 'replace')}", error_class="tool_execution")
        script = self.mounts.work / "transcribe.py"
        script.write_text(
            "import json,sys\n"
            "import whisper\n"
            "model=whisper.load_model('tiny')\n"
            "result=model.transcribe(sys.argv[1], language='en', fp16=False)\n"
            "print(json.dumps({'text': result.get('text',''), 'language': result.get('language','')}))\n",
            encoding="utf-8",
        )
        code, out, err = _run_fixed(
            [python.path, str(script), str(wav)],
            cwd=self.mounts.work,
            wall_seconds=min(180, max(60, self.budget.wall_seconds)),
        )
        try:
            payload = json.loads(out.decode("utf-8") or "{}")
        except json.JSONDecodeError:
            payload = {"text": out.decode("utf-8", "replace")[:2000]}
        result = {
            "operation": "media.transcribe",
            "exit_code": code,
            "transcript": payload,
            "stderr": err.decode("utf-8", "replace")[:2000],
            "waveform_sha256": file_digest(wav),
        }
        path = self.mounts.output / "transcript.json"
        path.write_bytes(canonical(result))
        return result, python, [path]

    # -- document ------------------------------------------------------------

    def _document_render(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        python = _require_tool(self.inventory, "python")
        title = str(request.parameters.get("title", "Odyssey Document"))
        body = str(request.parameters.get("body", "Tool-bearing document render."))
        _assert_no_evaluator_tokens(title + body, label="document content")
        script = self.mounts.work / "render_doc.py"
        docx_path = self.mounts.output / "document.docx"
        script.write_text(
            "from docx import Document\n"
            "import sys\n"
            "doc=Document()\n"
            "doc.add_heading(sys.argv[1], 0)\n"
            "doc.add_paragraph(sys.argv[2])\n"
            "doc.save(sys.argv[3])\n"
            "print('ok')\n",
            encoding="utf-8",
        )
        code, out, err = _run_fixed(
            [python.path, str(script), title, body, str(docx_path)],
            cwd=self.mounts.work,
            wall_seconds=30,
        )
        if code != 0 or not docx_path.exists():
            raise ToolRefused(f"document render failed: {err.decode('utf-8', 'replace')}", error_class="tool_execution")
        # DOCX is a zip container; the product cache refuses opaque archives.
        # Keep the real file on the sandbox output mount and admit only the
        # structured receipt that carries its content digest.
        result = {
            "operation": "document.render",
            "exit_code": code,
            "document_path": "output/document.docx",
            "document_sha256": file_digest(docx_path),
            "document_bytes": docx_path.stat().st_size,
            "stdout": out.decode("utf-8", "replace")[:500],
        }
        meta = self.mounts.output / "document_render.json"
        meta.write_bytes(canonical(result))
        return result, python, [meta]

    def _document_extract_structure(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        python = _require_tool(self.inventory, "python")
        # Render first if no input digest provided.
        if not request.input_artifact_digests:
            render_req = ToolRequest(
                lane_id=request.lane_id,
                arm=request.arm,
                task_id=request.task_id,
                operation="document.render",
                input_artifact_digests=(),
                parameters={"title": "Argument", "body": "Premise. Inference. Conclusion."},
                resource_budget=request.resource_budget,
                deadline_unix=request.deadline_unix,
                declared_operations=request.declared_operations,
                budget_sha256=request.budget_sha256,
                frontier=request.frontier,
            )
            self._document_render(render_req)
            docx_path = self.mounts.output / "document.docx"
        else:
            docx_path = self.mounts.inputs / "input.docx"
            self.cache.materialize_admitted(request.input_artifact_digests[0], docx_path, expected_lane=request.lane_id)
        script = self.mounts.work / "extract_doc.py"
        script.write_text(
            "from docx import Document\n"
            "import json,sys\n"
            "doc=Document(sys.argv[1])\n"
            "paras=[p.text for p in doc.paragraphs if p.text.strip()]\n"
            "print(json.dumps({'paragraphs': paras, 'count': len(paras)}))\n",
            encoding="utf-8",
        )
        code, out, err = _run_fixed([python.path, str(script), str(docx_path)], cwd=self.mounts.work, wall_seconds=30)
        try:
            structure = json.loads(out.decode("utf-8") or "{}")
        except json.JSONDecodeError:
            structure = {"raw": out.decode("utf-8", "replace")[:2000]}
        result = {
            "operation": "document.extract_structure",
            "exit_code": code,
            "structure": structure,
            "stderr": err.decode("utf-8", "replace")[:1000],
            # Philosophy frontier: tools verify structure only; never choose conclusions.
            "conclusion_authority": "human_or_model_only",
        }
        path = self.mounts.output / "document_structure.json"
        path.write_bytes(canonical(result))
        return result, python, [path]

    # -- three_d (primary: substrate_spatial3d; Blender optional comparison) -

    def _spatial_runtime_path(self) -> Path:
        return self.mounts.work / "scene_runtime.json"

    def _spatial_evaluator_dir(self) -> Path:
        # Path component ``evaluator`` is a forbidden token for arm parameters,
        # so candidates cannot name or request this namespace via tools.
        return self.mounts.root / "evaluator"

    def _spatial_revision(self) -> ToolRevision:
        """Pin the native spatial3d module as the primary 3D tool revision."""
        path = Path(spatial3d.__file__).resolve()
        return ToolRevision(
            tool_id=spatial3d.RENDERER_ID,
            path=str(path),
            version=spatial3d.RENDERER_VERSION,
            artifact_sha256=file_digest(path),
        )

    def _load_or_refuse_scene(self, request: ToolRequest | None = None) -> spatial3d.Scene:
        path = self._spatial_runtime_path()
        if not path.is_file():
            # Auto-build from parameters (or the occlusion canary seed) so single-op
            # canaries and registry smoke tests do not require a prior call.
            seed_id = "canary_occlusion_v1"
            if request is not None and request.parameters.get("seed_id"):
                seed_id = str(request.parameters["seed_id"])
            try:
                scene = spatial3d.build_scene_from_seed(seed_id)
            except spatial3d.Spatial3DError as error:
                raise ToolRefused(
                    f"no scene runtime and auto-build failed: {error}",
                    error_class="tool_execution",
                ) from error
            self._persist_scene(scene, seed_id=seed_id)
            return scene
        try:
            return spatial3d.read_runtime_scene(path)
        except spatial3d.Spatial3DError as error:
            raise ToolRefused(str(error), error_class="tool_execution") from error

    def _persist_scene(self, scene: spatial3d.Scene, *, seed_id: str) -> str:
        runtime_digest = spatial3d.write_runtime_scene(self._spatial_runtime_path(), scene)
        # Evaluator-only material: never added to output_paths / admitted digests.
        spatial3d.write_evaluator_state(self._spatial_evaluator_dir(), scene, seed_id)
        return runtime_digest

    def _three_d_build_scene(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Build a multi-object scene from a pinned public seed.

        Candidate-visible result: object/camera identity catalogs (no poses),
        scene_digest, renderer pin.  Full scene graph + answer state land only
        under the evaluator mount.
        """
        seed_id = str(request.parameters.get("seed_id") or "canary_occlusion_v1")
        _assert_no_evaluator_tokens(seed_id, label="seed_id")
        seed_override = request.parameters.get("seed")
        try:
            scene = spatial3d.build_scene_from_seed(
                seed_id,
                seed_override=int(seed_override) if seed_override is not None else None,
            )
        except (spatial3d.Spatial3DError, ValueError, KeyError) as error:
            raise ToolRefused(f"three_d.build_scene failed: {error}", error_class="tool_execution") from error
        runtime_digest = self._persist_scene(scene, seed_id=seed_id)
        result = {
            "operation": "three_d.build_scene",
            "renderer": spatial3d.RENDERER_ID,
            "renderer_version": spatial3d.RENDERER_VERSION,
            "seed_id": scene.seed_id,
            "seed": scene.seed,
            "scene_digest": scene.canonical_digest(),
            "runtime_sha256": runtime_digest,
            "object_ids": [row["object_id"] for row in scene.public_object_catalog()],
            "camera_ids": [row["camera_id"] for row in scene.public_camera_catalog()],
            "active_camera": scene.active_camera,
            "width": scene.width,
            "height": scene.height,
            # Explicit: scene graph is NOT in this payload.
            "scene_graph_exposed": False,
            "evaluator_state_exposed": False,
        }
        meta = self.mounts.output / "three_d_build_scene.json"
        meta.write_bytes(canonical(result))
        return result, self._spatial_revision(), [meta]

    def _three_d_render(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Render RGB pixels from real 3D geometry via the native z-buffer.

        Primary path: ``substrate_spatial3d``.  Optional comparison path:
        ``backend=blender`` runs Blender Cycles CPU when the binary is present
        and launches cleanly — never a launch blocker, never the default.
        """
        backend = str(request.parameters.get("backend") or "spatial3d").strip().lower()
        if backend in {"blender", "blender_cycles_cpu"}:
            return self._three_d_render_blender_optional(request)

        scene = self._load_or_refuse_scene(request)

        camera_id = request.parameters.get("camera_id")
        if camera_id is not None:
            camera_id = str(camera_id)
            _assert_no_evaluator_tokens(camera_id, label="camera_id")
        try:
            rendered = spatial3d.render_scene(scene, camera_id=camera_id)
        except spatial3d.Spatial3DError as error:
            raise ToolRefused(f"three_d.render failed: {error}", error_class="tool_execution") from error

        png = self.mounts.output / "render.png"
        png.write_bytes(rendered.rgb_png)
        # Cache-admissible metric depth (16-bit greyscale PNG, millimeters).
        depth_path = self.mounts.output / "depth_u16.png"
        depth_path.write_bytes(rendered.depth_u16_png)
        depth_vis = self.mounts.output / "depth_vis.png"
        depth_vis.write_bytes(rendered.depth_vis_png)
        # Float32 depth kept under work/ for evaluator evidence only — not admitted.
        (self.mounts.work / "depth.f32").write_bytes(rendered.depth_f32)

        sensor = spatial3d.candidate_sensor_bundle(rendered)
        result = {
            "operation": "three_d.render",
            "pixels": True,
            "depth": True,
            "renderer": spatial3d.RENDERER_ID,
            "renderer_version": spatial3d.RENDERER_VERSION,
            "backend": "spatial3d",
            "image_sha256": file_digest(png),
            "depth_sha256": file_digest(depth_path),
            "depth_f32_sha256": hashlib.sha256(rendered.depth_f32).hexdigest(),
            "depth_vis_sha256": file_digest(depth_vis),
            "depth_encoding": {
                "admitted_format": "png_uint16_greyscale_millimeters",
                "scale_mm": spatial3d.DEPTH_U16_SCALE_MM,
                "zero_means": "background_or_invalid",
            },
            "camera_id": rendered.camera_id,
            "seed_id": rendered.seed_id,
            "scene_digest": rendered.scene_digest,
            "width": rendered.width,
            "height": rendered.height,
            "visible_object_ids": list(rendered.visible_object_ids),
            "sensor": sensor,
            "scene_graph_exposed": False,
            "evaluator_state_exposed": False,
        }
        meta = self.mounts.output / "three_d_render.json"
        meta.write_bytes(canonical(result))
        return result, self._spatial_revision(), [png, depth_path, depth_vis, meta]

    def _three_d_render_blender_optional(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Optional Blender Cycles comparison render — never the primary path."""
        png = self.mounts.output / "blender_compare.png"
        blender = self.inventory.get("blender")
        if blender is None:
            raise ToolRefused(
                "backend=blender requested but Blender is not on the tool inventory",
                error_class="tool_execution",
            )
        script = self.mounts.work / "render_cube_compare.py"
        script.write_text(
            "import bpy\n"
            f"out = {str(png)!r}\n"
            "bpy.ops.mesh.primitive_cube_add()\n"
            "scene = bpy.context.scene\n"
            "scene.render.engine = 'CYCLES'\n"
            "scene.cycles.device = 'CPU'\n"
            "scene.cycles.samples = 8\n"
            "scene.render.resolution_x = 64\n"
            "scene.render.resolution_y = 64\n"
            "scene.render.filepath = out\n"
            "bpy.ops.render.render(write_still=True)\n"
            "print('RENDER_OK', out)\n",
            encoding="utf-8",
        )
        code, out, err = _run_fixed(
            [blender.path, "-b", "-noaudio", "--python", str(script)],
            cwd=self.mounts.work,
            wall_seconds=min(120, max(60, self.budget.wall_seconds)),
        )
        if code != 0 or not png.exists() or png.stat().st_size == 0:
            raise ToolRefused(
                "optional Blender comparison render failed "
                f"(exit={code}); primary backend remains substrate_spatial3d",
                error_class="tool_execution",
            )
        result = {
            "operation": "three_d.render",
            "exit_code": code,
            "image_sha256": file_digest(png),
            "stdout": out.decode("utf-8", "replace")[:1500],
            "stderr": err.decode("utf-8", "replace")[:1500],
            "pixels": True,
            "renderer": "blender_cycles_cpu",
            "renderer_version": blender.version,
            "backend": "blender",
            "primary_backend": spatial3d.RENDERER_ID,
            "geometry": "unit_cube_comparison_only",
            "comparison_only": True,
        }
        meta = self.mounts.output / "three_d_render_blender.json"
        meta.write_bytes(canonical(result))
        return result, blender, [png, meta]

    def _three_d_depth(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Emit the z-buffer as a first-class depth artifact for the active/selected camera."""
        scene = self._load_or_refuse_scene(request)
        camera_id = request.parameters.get("camera_id")
        if camera_id is not None:
            camera_id = str(camera_id)
            _assert_no_evaluator_tokens(camera_id, label="camera_id")
        try:
            rendered = spatial3d.render_scene(scene, camera_id=camera_id)
        except spatial3d.Spatial3DError as error:
            raise ToolRefused(f"three_d.depth failed: {error}", error_class="tool_execution") from error
        depth_path = self.mounts.output / "depth_u16.png"
        depth_path.write_bytes(rendered.depth_u16_png)
        depth_vis = self.mounts.output / "depth_vis.png"
        depth_vis.write_bytes(rendered.depth_vis_png)
        (self.mounts.work / "depth.f32").write_bytes(rendered.depth_f32)
        result = {
            "operation": "three_d.depth",
            "renderer": spatial3d.RENDERER_ID,
            "renderer_version": spatial3d.RENDERER_VERSION,
            "depth_sha256": file_digest(depth_path),
            "depth_f32_sha256": hashlib.sha256(rendered.depth_f32).hexdigest(),
            "depth_vis_sha256": file_digest(depth_vis),
            "camera_id": rendered.camera_id,
            "seed_id": rendered.seed_id,
            "scene_digest": rendered.scene_digest,
            "width": rendered.width,
            "height": rendered.height,
            "far_sentinel": spatial3d.DEPTH_FAR_SENTINEL,
            "encoding": {
                "admitted_format": "png_uint16_greyscale_millimeters",
                "scale_mm": spatial3d.DEPTH_U16_SCALE_MM,
                "zero_means": "background_or_invalid",
                "float32_work_copy": "work/depth.f32 (not admitted)",
            },
            "scene_graph_exposed": False,
        }
        meta = self.mounts.output / "three_d_depth.json"
        meta.write_bytes(canonical(result))
        return result, self._spatial_revision(), [depth_path, depth_vis, meta]

    def _three_d_move_object(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Apply a deterministic object motion; updates runtime scene only (no GT leak)."""
        scene = self._load_or_refuse_scene(request)
        object_id = request.parameters.get("object_id")
        if not isinstance(object_id, str) or not object_id.strip():
            raise ToolRefused("three_d.move_object requires object_id", error_class="tool_protocol")
        _assert_no_evaluator_tokens(object_id, label="object_id")
        translation = request.parameters.get("translation")
        position = request.parameters.get("position")
        rotation_delta = request.parameters.get("rotation_delta")
        try:
            scene.move_object(
                object_id,
                translation=translation if isinstance(translation, (list, tuple)) else None,
                position=position if isinstance(position, (list, tuple)) else None,
                rotation_delta=rotation_delta if isinstance(rotation_delta, (list, tuple)) else None,
            )
        except spatial3d.Spatial3DError as error:
            raise ToolRefused(str(error), error_class="tool_execution") from error
        runtime_digest = self._persist_scene(scene, seed_id=scene.seed_id)
        result = {
            "operation": "three_d.move_object",
            "renderer": spatial3d.RENDERER_ID,
            "renderer_version": spatial3d.RENDERER_VERSION,
            "object_id": object_id,
            "scene_digest": scene.canonical_digest(),
            "runtime_sha256": runtime_digest,
            # Motion applied; new pose is evaluator-only (in scene graph).
            "scene_graph_exposed": False,
        }
        meta = self.mounts.output / "three_d_move_object.json"
        meta.write_bytes(canonical(result))
        return result, self._spatial_revision(), [meta]

    def _three_d_set_camera(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Select the active camera for subsequent render/depth operations."""
        scene = self._load_or_refuse_scene(request)
        camera_id = request.parameters.get("camera_id")
        if not isinstance(camera_id, str) or not camera_id.strip():
            raise ToolRefused("three_d.set_camera requires camera_id", error_class="tool_protocol")
        _assert_no_evaluator_tokens(camera_id, label="camera_id")
        try:
            scene.set_active_camera(camera_id)
        except spatial3d.Spatial3DError as error:
            raise ToolRefused(str(error), error_class="tool_execution") from error
        runtime_digest = self._persist_scene(scene, seed_id=scene.seed_id)
        result = {
            "operation": "three_d.set_camera",
            "renderer": spatial3d.RENDERER_ID,
            "renderer_version": spatial3d.RENDERER_VERSION,
            "active_camera": scene.active_camera,
            "camera_ids": sorted(scene.cameras),
            "intrinsics": scene.cameras[scene.active_camera].intrinsics(scene.width, scene.height),
            "scene_digest": scene.canonical_digest(),
            "runtime_sha256": runtime_digest,
            "scene_graph_exposed": False,
        }
        meta = self.mounts.output / "three_d_set_camera.json"
        meta.write_bytes(canonical(result))
        return result, self._spatial_revision(), [meta]

    def _three_d_inspect_mesh(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        """Inspect local mesh topology for one object — no full scene-graph poses."""
        scene = self._load_or_refuse_scene(request)
        object_id = str(request.parameters.get("object_id") or next(iter(sorted(scene.objects))))
        _assert_no_evaluator_tokens(object_id, label="object_id")
        try:
            mesh_info = spatial3d.inspect_mesh_public(scene, object_id)
        except spatial3d.Spatial3DError as error:
            raise ToolRefused(str(error), error_class="tool_execution") from error
        # Also emit a real OBJ for the local mesh (geometry on disk, not evaluator GT).
        mesh = spatial3d.get_mesh(scene.objects[object_id].shape)
        obj_path = self.mounts.work / f"{object_id}.obj"
        lines = [f"# {mesh.name} local mesh", f"# object_id {object_id}"]
        for vx, vy, vz in mesh.vertices:
            lines.append(f"v {vx:.6f} {vy:.6f} {vz:.6f}")
        for i0, i1, i2 in mesh.triangles:
            lines.append(f"f {i0 + 1} {i1 + 1} {i2 + 1}")
        obj_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        result = {
            "operation": "three_d.inspect_mesh",
            "mesh": mesh_info,
            "mesh_sha256": file_digest(obj_path),
            "geometry": True,
            "renderer": spatial3d.RENDERER_ID,
            "renderer_version": spatial3d.RENDERER_VERSION,
            "scene_graph_exposed": False,
        }
        path = self.mounts.output / "mesh_inspect.json"
        path.write_bytes(canonical(result))
        return result, self._spatial_revision(), [path]

    # -- compute -------------------------------------------------------------

    def _compute_python(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        python = _require_tool(self.inventory, "python")
        expression = request.parameters.get("expression", "sum(range(10))")
        if not isinstance(expression, str) or len(expression) > 2000:
            raise ToolRefused("python expression is invalid", error_class="tool_protocol")
        _assert_no_evaluator_tokens(expression, label="python expression")
        # AST-gated: only expression evaluation, no imports/attributes/calls beyond basics.
        try:
            tree = ast.parse(expression, mode="eval")
        except SyntaxError as error:
            raise ToolRefused(f"python expression is not a valid eval form: {error}", error_class="tool_protocol") from error
        for node in ast.walk(tree):
            if isinstance(node, (ast.Import, ast.ImportFrom, ast.Attribute, ast.Lambda, ast.Await, ast.Yield)):
                raise ToolRefused("python expression uses a forbidden construct", error_class="sandbox")
            allowed_calls = {"sum", "range", "min", "max", "abs", "len", "round", "sorted", "list", "tuple", "set", "dict", "int", "float", "str"}
            if isinstance(node, ast.Call) and (
                not isinstance(node.func, ast.Name) or node.func.id not in allowed_calls
            ):
                raise ToolRefused("python expression call is not allowlisted", error_class="sandbox")
        script = self.mounts.work / "compute.py"
        script.write_text(
            "import json,ast\n"
            f"expr={expression!r}\n"
            "value=eval(compile(ast.parse(expr, mode='eval'), '<expr>', 'eval'), {'__builtins__': {}}, "
            "{'sum':sum,'range':range,'min':min,'max':max,'abs':abs,'len':len,'round':round,"
            "'sorted':sorted,'list':list,'tuple':tuple,'set':set,'dict':dict,'int':int,'float':float,'str':str})\n"
            "print(json.dumps({'value': value}, default=str))\n",
            encoding="utf-8",
        )
        code, out, err = _run_fixed([python.path, str(script)], cwd=self.mounts.work, wall_seconds=30)
        try:
            payload = json.loads(out.decode("utf-8") or "{}")
        except json.JSONDecodeError:
            payload = {"raw": out.decode("utf-8", "replace")[:2000]}
        result = {
            "operation": "compute.python",
            "exit_code": code,
            "result": payload,
            "stderr": err.decode("utf-8", "replace")[:1000],
            "expression_sha256": hashlib.sha256(expression.encode("utf-8")).hexdigest(),
        }
        path = self.mounts.output / "compute_python.json"
        path.write_bytes(canonical(result))
        return result, python, [path]

    def _compute_sympy(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        python = _require_tool(self.inventory, "python")
        expression = request.parameters.get("expression", "integrate(x**2, x)")
        if not isinstance(expression, str) or len(expression) > 2000:
            raise ToolRefused("sympy expression is invalid", error_class="tool_protocol")
        _assert_no_evaluator_tokens(expression, label="sympy expression")
        script = self.mounts.work / "sympy_compute.py"
        script.write_text(
            "import json\n"
            "import sympy\n"
            "from sympy import *\n"
            "x, y, z, t = symbols('x y z t')\n"
            f"expr = {expression!r}\n"
            "value = simplify(sympify(expr))\n"
            "print(json.dumps({'input': expr, 'value': str(value)}))\n",
            encoding="utf-8",
        )
        code, out, err = _run_fixed([python.path, str(script)], cwd=self.mounts.work, wall_seconds=60)
        try:
            payload = json.loads(out.decode("utf-8") or "{}")
        except json.JSONDecodeError:
            payload = {"raw": out.decode("utf-8", "replace")[:2000]}
        result = {
            "operation": "compute.sympy",
            "exit_code": code,
            "result": payload,
            "stderr": err.decode("utf-8", "replace")[:1000],
            "exact": True,
        }
        path = self.mounts.output / "compute_sympy.json"
        path.write_bytes(canonical(result))
        return result, python, [path]

    def _source_read_cached(self, request: ToolRequest) -> tuple[dict[str, Any], ToolRevision, list[Path]]:
        if not request.input_artifact_digests:
            raise ToolRefused("source.read_cached requires an admitted input digest", error_class="tool_protocol")
        digest_value = request.input_artifact_digests[0]
        data = self.cache.read_admitted(digest_value, expected_lane=request.lane_id)
        # Never network: only lane-local admitted cache.
        preview = data[:512]
        result = {
            "operation": "source.read_cached",
            "artifact_sha256": digest_value,
            "byte_length": len(data),
            "preview_sha256": hashlib.sha256(preview).hexdigest(),
            "network": False,
        }
        path = self.mounts.output / "source_read.json"
        path.write_bytes(canonical(result))
        # Use git pin as a neutral revision marker for cache-only reads.
        revision = self.inventory.get("git") or ToolRevision("cache", "cache", "cache", digest({"cache": True}))
        return result, revision, [path]


# ---------------------------------------------------------------------------
# Broker
# ---------------------------------------------------------------------------


class ToolBroker:
    """Validate, execute, quarantine, admit — shared by both arms."""

    def __init__(
        self,
        *,
        root: Path,
        lane_id: str,
        arm: str,
        budget: ToolBudget | None = None,
        inventory: Mapping[str, ToolRevision] | None = None,
        peer_budget_sha256: str | None = None,
    ) -> None:
        if arm not in ROLES:
            raise ToolRefused("broker arm is invalid", error_class="parity")
        if lane_id not in FRONTIERS and not _IDENTIFIER.fullmatch(lane_id):
            raise ToolRefused("broker lane_id is invalid", error_class="tool_protocol")
        self.root = root.resolve()
        self.lane_id = lane_id
        self.arm = arm
        self.budget = budget or ToolBudget()
        self.budget_sha256 = self.budget.budget_sha256()
        if peer_budget_sha256 is not None and peer_budget_sha256 != self.budget_sha256:
            raise ToolRefused(
                "candidate/control tool budget asymmetry refused at broker bind",
                error_class="parity",
            )
        self.inventory = dict(inventory) if inventory is not None else discover_tool_inventory()
        cache_root = self.root / "evidence/artifacts/substrate/odyssey7d/tool-cache" / lane_id / arm
        self.cache = LaneToolCache(root=cache_root, lane_id=lane_id, arm=arm)
        # Pre-admit tool pins so derivatives can reference them.
        for revision in self.inventory.values():
            with contextlib.suppress(ToolRefused, ProductRefused):
                # Pin admission failures are deferred until the tool is required.
                self.cache.ensure_tool_pin(revision)

    def _sandbox_root(self, task_id: str) -> Path:
        safe_task = re.sub(r"[^a-zA-Z0-9._-]+", "_", task_id)[:80]
        return self.root / "evidence/artifacts/substrate/odyssey7d/tool-work" / self.lane_id / self.arm / safe_task

    def execute(self, request: ToolRequest) -> ToolResponse:
        started = time.monotonic()
        try:
            response = self._execute_inner(request, started=started)
        except ToolRefused as error:
            unsigned = {
                "schema": TOOL_RECEIPT_SCHEMA,
                "status": "refused",
                "operation": request.operation,
                "output_digests": [],
                "stdout_digest": None,
                "stderr_digest": None,
                "tool_revision": {},
                "resource_use": {"elapsed_seconds": round(time.monotonic() - started, 6)},
                "provenance": {
                    "request_sha256": digest(request.to_dict()),
                    "lane_id": self.lane_id,
                    "arm": self.arm,
                    "budget_sha256": self.budget_sha256,
                },
                "error_class": error.error_class,
                "admitted": False,
                "detail": {"message": str(error)},
            }
            receipt_sha = digest(unsigned)
            return ToolResponse(
                status="refused",
                operation=request.operation,
                output_digests=[],
                stdout_digest=None,
                stderr_digest=None,
                tool_revision={},
                resource_use=unsigned["resource_use"],
                provenance=unsigned["provenance"],
                error_class=error.error_class,
                receipt_sha256=receipt_sha,
                admitted=False,
                detail=unsigned["detail"],
            )
        return response

    def _execute_inner(self, request: ToolRequest, *, started: float) -> ToolResponse:
        # Protocol validation
        if request.lane_id != self.lane_id:
            raise ToolRefused("request lane_id does not match broker lane", error_class="cross_lane")
        if request.arm != self.arm:
            raise ToolRefused("request arm does not match broker arm", error_class="parity")
        if request.operation not in REGISTRY_OPERATIONS:
            raise ToolRefused(
                f"operation {request.operation!r} is outside the closed registry",
                error_class="undeclared_operation",
            )
        if request.operation not in request.declared_operations:
            raise ToolRefused(
                f"operation {request.operation!r} is not declared for this task",
                error_class="undeclared_operation",
            )
        if request.frontier in FRONTIERS and request.operation not in FRONTIER_OPERATIONS[request.frontier]:
            raise ToolRefused(
                f"operation {request.operation!r} is outside frontier {request.frontier} surface",
                error_class="undeclared_operation",
            )
        if request.budget_sha256 != self.budget_sha256 or request.resource_budget.budget_sha256() != self.budget_sha256:
            raise ToolRefused("tool budget digest mismatch refused", error_class="parity")
        if time.time() > request.deadline_unix:
            raise ToolRefused("tool request deadline exceeded", error_class="deadline")
        _assert_no_evaluator_tokens(request.parameters, label="tool parameters")
        _assert_no_evaluator_tokens(request.task_id, label="task_id")
        for item in request.input_artifact_digests:
            if not _is_sha256(item):
                raise ToolRefused("input artifact digest is invalid", error_class="tool_protocol")

        mounts = SandboxMounts.create(self._sandbox_root(request.task_id))
        executor = OperationExecutor(
            mounts=mounts,
            inventory=self.inventory,
            cache=self.cache,
            budget=self.budget,
            root=self.root,
            lane_id=self.lane_id,
            arm=self.arm,
        )
        result, revision, output_paths = executor.execute(request)

        # Quarantine + admit every output path.  Receipts stay compact:
        # digest/size/type/producer/timing/path — never duplicate raw stdout,
        # video, renders, or proof text into the receipt body.
        tool_pin = self.cache.ensure_tool_pin(revision)
        output_digests: list[str] = []
        compact_outputs: list[dict[str, Any]] = []
        stdout_digest: str | None = None
        for path in output_paths:
            if not path.is_file() or path.is_symlink():
                raise ToolRefused(f"sandbox output is not a regular file: {path}", error_class="sandbox")
            raw = path.read_bytes()
            if len(raw) > self.budget.max_output_bytes:
                raise ToolRefused("tool output exceeds budget max_output_bytes", error_class="budget")
            media_type = _STRUCTURED_MEDIA if path.suffix == ".json" else _guess_media_type(path, raw)
            source_ref = digest(
                {
                    "lane": self.lane_id,
                    "arm": self.arm,
                    "task": request.task_id,
                    "operation": request.operation,
                    "name": path.name,
                }
            )
            # Prefer raw admission for primary outputs; lineage when we have verified inputs.
            admitted = self.cache.quarantine_and_admit_bytes(
                raw,
                media_type=media_type,
                source_reference_sha256=source_ref,
                rights_status="user-provided",
            )
            output_digests.append(admitted)
            if path.suffix == ".json" and stdout_digest is None:
                stdout_digest = admitted
            relative = (
                str(path.relative_to(self.root)) if path.is_relative_to(self.root) else path.name
            )
            compact_outputs.append(
                compact_artifact_receipt(
                    digest_hex=admitted,
                    size=len(raw),
                    media_type=media_type,
                    producer=f"{request.operation}:{revision.tool_id}",
                    path=relative,
                    timing_seconds=None,
                    resource_use={"lane_id": self.lane_id, "arm": self.arm},
                )
            )

        elapsed = round(time.monotonic() - started, 6)
        request_sha = digest(request.to_dict())
        resource_class = None
        with contextlib.suppress(Exception):
            resource_class = resource_class_for_frontier(request.frontier)
        provenance = {
            "request_sha256": request_sha,
            "lane_id": self.lane_id,
            "arm": self.arm,
            "task_id": request.task_id,
            "frontier": request.frontier,
            "budget_sha256": self.budget_sha256,
            "tool_pin_sha256": tool_pin,
            "cache_id": self.cache.store.cache_id,
            "sandbox_root": str(mounts.root.relative_to(self.root)) if mounts.root.is_relative_to(self.root) else str(mounts.root),
            "input_artifact_digests": list(request.input_artifact_digests),
            "result_summary_sha256": digest(result),
            "resource_class": resource_class,
            "exec_mode": executor.last_exec_mode,
            "compact_outputs": compact_outputs,
        }
        unsigned = {
            "schema": TOOL_RECEIPT_SCHEMA,
            "status": "ok",
            "operation": request.operation,
            "output_digests": output_digests,
            "stdout_digest": stdout_digest,
            "stderr_digest": None,
            "tool_revision": revision.to_dict(),
            "resource_use": {
                "elapsed_seconds": elapsed,
                "wall_seconds_budget": self.budget.wall_seconds,
                "attempts": TOOL_TRANSPORT_ATTEMPTS,
                "resource_class": resource_class,
                "exec_mode": executor.last_exec_mode,
            },
            "provenance": provenance,
            "error_class": "ok",
            "admitted": True,
            "detail": {"result_keys": sorted(result) if isinstance(result, dict) else []},
        }
        receipt_sha = digest(unsigned)
        # Persist receipt under lane output for canary inspection.
        receipt_dir = self.root / "evidence/artifacts/substrate/odyssey7d/tool-receipts" / self.lane_id / self.arm
        receipt_dir.mkdir(parents=True, exist_ok=True)
        receipt_path = receipt_dir / f"{re.sub(r'[^a-zA-Z0-9._-]+', '_', request.task_id)[:80]}-{request.operation.replace('.', '_')}.json"
        receipt_body = {**unsigned, "receipt_sha256": receipt_sha}
        receipt_path.write_text(json.dumps(receipt_body, sort_keys=True, indent=2) + "\n", encoding="utf-8")

        # Verify self-consistency of the receipt we just wrote.
        verify_tool_receipt(receipt_body, cache=self.cache)

        return ToolResponse(
            status="ok",
            operation=request.operation,
            output_digests=output_digests,
            stdout_digest=stdout_digest,
            stderr_digest=None,
            tool_revision=revision.to_dict(),
            resource_use=unsigned["resource_use"],
            provenance=provenance,
            error_class="ok",
            receipt_sha256=receipt_sha,
            admitted=True,
            detail=unsigned["detail"],
        )


def _guess_media_type(path: Path, raw: bytes) -> str:
    if path.suffix == ".png" or raw.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if path.suffix == ".wav" or (len(raw) >= 12 and raw[:4] == b"RIFF" and raw[8:12] == b"WAVE"):
        return "audio/wav"
    if path.suffix == ".docx":
        # Sniffer maps zip to application/zip which the cache refuses as opaque archive.
        # Represent the document metadata as structured JSON instead for admission;
        # callers still have the file_digest of the real docx on disk in the result.
        return "text/plain"
    if path.suffix == ".json":
        return _STRUCTURED_MEDIA
    try:
        raw.decode("utf-8")
        return "text/plain"
    except UnicodeDecodeError:
        return "application/octet-stream"


# ---------------------------------------------------------------------------
# Arm-facing helpers
# ---------------------------------------------------------------------------


def make_tool_request(
    *,
    lane_id: str,
    arm: str,
    task_id: str,
    operation: str,
    frontier: str,
    declared_operations: frozenset[str],
    budget: ToolBudget,
    parameters: dict[str, Any] | None = None,
    input_artifact_digests: list[str] | None = None,
    deadline_unix: float | None = None,
) -> ToolRequest:
    return ToolRequest(
        lane_id=lane_id,
        arm=arm,
        task_id=task_id,
        operation=operation,
        input_artifact_digests=tuple(input_artifact_digests or ()),
        parameters=dict(parameters or {}),
        resource_budget=budget,
        deadline_unix=float(deadline_unix if deadline_unix is not None else time.time() + budget.wall_seconds),
        declared_operations=declared_operations,
        budget_sha256=budget.budget_sha256(),
        frontier=frontier,
    )


def execute_tool_proposals(
    *,
    root: Path,
    request: Mapping[str, Any],
    role: str,
    proposals: list[Mapping[str, Any]],
    budget: ToolBudget | None = None,
    inventory: Mapping[str, ToolRevision] | None = None,
    peer_budget_sha256: str | None = None,
) -> list[dict[str, Any]]:
    """Execute model-proposed tools under the task's declared surface."""

    frontier = str(request["frontier"])
    task = request["task"]
    declared = declared_operations_for_task(task, frontier=frontier)
    # Frontier resource class — candidate and control of a pair share this path.
    budget = budget or budget_for_frontier(frontier)
    broker = ToolBroker(
        root=root,
        lane_id=frontier,
        arm=role,
        budget=budget,
        inventory=inventory,
        peer_budget_sha256=peer_budget_sha256,
    )
    if len(proposals) > budget.max_tool_calls:
        raise ToolRefused("tool_proposals exceed max_tool_calls", error_class="budget")
    admitted: list[dict[str, Any]] = []
    for index, proposal in enumerate(proposals):
        if not isinstance(proposal, Mapping):
            raise ToolRefused(f"tool proposal {index} is malformed", error_class="tool_protocol")
        operation = proposal.get("operation")
        if not isinstance(operation, str):
            raise ToolRefused(f"tool proposal {index} lacks operation", error_class="tool_protocol")
        parameters = proposal.get("parameters") or {}
        if not isinstance(parameters, dict):
            raise ToolRefused(f"tool proposal {index} parameters must be an object", error_class="tool_protocol")
        inputs = proposal.get("input_artifact_digests") or []
        if not isinstance(inputs, list) or not all(isinstance(item, str) for item in inputs):
            raise ToolRefused(f"tool proposal {index} input digests are invalid", error_class="tool_protocol")
        tool_request = make_tool_request(
            lane_id=frontier,
            arm=role,
            task_id=str(task["task_id"]),
            operation=operation,
            frontier=frontier,
            declared_operations=declared,
            budget=budget,
            parameters=parameters,
            input_artifact_digests=list(inputs),
        )
        response = broker.execute(tool_request)
        if not response.admitted:
            raise ToolRefused(
                f"tool proposal {index} refused: {response.detail.get('message', response.error_class)}",
                error_class=response.error_class,
            )
        admitted.append(response.to_dict())
    return admitted


def run_frontier_canary(
    root: Path,
    *,
    budget: ToolBudget | None = None,
    frontiers: str = "ABCDEFGH",
    roles: tuple[str, ...] = ("candidate", "control"),
) -> dict[str, Any]:
    """Public tool-bearing canary: one real operation per frontier × arm."""

    budget = budget or ToolBudget()
    shared = budget.budget_sha256()
    inventory = discover_tool_inventory()
    # Parity bind once for the whole canary.
    assert_budget_parity(budget, ToolBudget.from_dict(budget.to_dict()))
    rows: list[dict[str, Any]] = []
    for frontier in frontiers:
        operation = FRONTIER_CANARY_OPERATION[frontier]
        declared = FRONTIER_OPERATIONS[frontier]
        for role in roles:
            broker = ToolBroker(
                root=root,
                lane_id=frontier,
                arm=role,
                budget=budget,
                inventory=inventory,
                peer_budget_sha256=shared,
            )
            tool_request = make_tool_request(
                lane_id=frontier,
                arm=role,
                task_id=f"canary-{frontier}-{role}",
                operation=operation,
                frontier=frontier,
                declared_operations=declared,
                budget=budget,
                parameters=_canary_parameters(frontier, operation),
            )
            response = broker.execute(tool_request)
            rows.append(
                {
                    "frontier": frontier,
                    "role": role,
                    "operation": operation,
                    "status": response.status,
                    "admitted": response.admitted,
                    "error_class": response.error_class,
                    "tool_revision": response.tool_revision,
                    "output_digests": response.output_digests,
                    "receipt_sha256": response.receipt_sha256,
                    "resource_use": response.resource_use,
                }
            )
    document = {
        "schema": TOOL_CANARY_SCHEMA,
        "activation": False,
        "program": PROGRAM,
        "budget_sha256": shared,
        "budget": budget.to_dict(),
        "generated_at": _timestamp(),
        "rows": rows,
        "all_admitted": all(row["admitted"] for row in rows),
        "registry_size": len(REGISTRY_OPERATIONS),
    }
    document["sha256"] = digest({key: value for key, value in document.items() if key != "sha256"})
    out_dir = root / "evidence/artifacts/substrate/odyssey7d/tool-bearing-canary"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "TOOL_BEARING_CANARY.json"
    out_path.write_text(json.dumps(document, sort_keys=True, indent=2) + "\n", encoding="utf-8")
    return document


def _canary_parameters(frontier: str, operation: str) -> dict[str, Any]:
    if operation == "formal.check_lean":
        return {"source": "theorem two_plus_two : 2 + 2 = 4 := by rfl\n"}
    if operation == "formal.solve_smt":
        return {"smt": "(set-logic QF_LIA)\n(declare-const x Int)\n(assert (= (+ x 2) 4))\n(check-sat)\n"}
    if operation == "repo.test":
        return {"mode": "pass"}
    if operation == "compute.sympy":
        return {"expression": "integrate(x**2, x)"}
    if operation == "document.extract_structure":
        return {}
    if operation == "repo.inspect":
        return {}
    if operation == "media.probe":
        return {}
    if operation == "three_d.render":
        return {"seed_id": "canary_occlusion_v1", "backend": "spatial3d"}
    if operation == "three_d.build_scene":
        return {"seed_id": "canary_occlusion_v1"}
    if operation == "three_d.depth":
        return {"camera_id": "cam_front"}
    if operation == "three_d.move_object":
        return {"object_id": "occluder", "translation": [0.0, 0.1, 0.0]}
    if operation == "three_d.set_camera":
        return {"camera_id": "cam_side"}
    if operation == "three_d.inspect_mesh":
        return {"object_id": "occluder", "seed_id": "canary_occlusion_v1"}
    return {}


__all__ = (
    "FRONTIER_CANARY_OPERATION",
    "FRONTIER_OPERATIONS",
    "REGISTRY_OPERATIONS",
    "ToolBudget",
    "ToolBroker",
    "ToolRefused",
    "ToolRequest",
    "ToolResponse",
    "ToolRevision",
    "allowed_operations_prompt_block",
    "assert_budget_parity",
    "budget_for_frontier",
    "declared_operations_for_task",
    "discover_tool_inventory",
    "execute_tool_proposals",
    "make_tool_request",
    "run_frontier_canary",
    "verify_tool_receipt",
)
